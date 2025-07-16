import tkinter as tk
from tkinter import filedialog, ttk, messagebox
import threading
import requests
from openpyxl import load_workbook
from docx import Document
from docx.shared import RGBColor
import time

# Global flags
paused = False
stopped = False

class RedirectCheckerApp:
    def __init__(self, root):
        self.root = root
        self.root.title("URL Redirect Checker")
        self.create_widgets()

    def create_widgets(self):
        self.select_button = tk.Button(self.root, text="Select Excel File", command=self.select_file)
        self.select_button.pack(pady=10)

        self.progress = ttk.Progressbar(self.root, length=400, mode='determinate')
        self.progress.pack(pady=5)

        self.pause_button = tk.Button(self.root, text="Pause", command=self.pause_resume)
        self.pause_button.pack(pady=5)

        self.stop_button = tk.Button(self.root, text="Stop", command=self.confirm_stop)
        self.stop_button.pack(pady=5)

        self.start_button = tk.Button(self.root, text="Start Checking", command=self.start_checking)
        self.start_button.pack(pady=10)

        self.status_label = tk.Label(self.root, text="Status: Waiting...")
        self.status_label.pack()

    def select_file(self):
        self.filepath = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx")])
        if self.filepath:
            messagebox.showinfo("File Selected", f"Selected: {self.filepath}")

    def pause_resume(self):
        global paused
        paused = not paused
        self.pause_button.config(text="Resume" if paused else "Pause")

    def confirm_stop(self):
        global stopped
        result = messagebox.askyesno("Confirm Stop", "Are you sure you want to stop the process?")
        if result:
            stopped = True

    def start_checking(self):
        global stopped
        stopped = False
        thread = threading.Thread(target=self.check_redirects)
        thread.start()

    def check_redirects(self):
        global paused, stopped

        if not hasattr(self, 'filepath') or not self.filepath:
            messagebox.showwarning("No File", "Please select an Excel file first.")
            return

        try:
            wb = load_workbook(self.filepath)
            ws = wb.active
            total = ws.max_row - 1
            self.progress["maximum"] = total

            document = Document()
            document.add_heading('URL Redirect Report', 0)

            table = document.add_table(rows=1, cols=4)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '404 Page'
            hdr_cells[1].text = 'Expected Redirect'
            hdr_cells[2].text = 'Final URL'
            hdr_cells[3].text = 'Status'

            for i, row in enumerate(ws.iter_rows(min_row=2, max_col=2), start=2):
                while paused:
                    self.status_label.config(text=f"Status: Paused at row {i}")
                    time.sleep(1)
                if stopped:
                    self.status_label.config(text="Stopped by user. Saving partial report...")
                    break

                source = row[0].value
                target = row[1].value
                if not source or not target:
                    continue  # Skip empty rows

                final_url = ""

                try:
                    resp = requests.head(source, allow_redirects=True, timeout=5)
                    final_url = resp.url

                    if final_url.strip('/') == target.strip('/'):
                        status = "Success"
                    else:
                        status = f"Mismatch"
                except Exception as e:
                    status = f"Error: {str(e)}"

                cells = table.add_row().cells
                cells[0].text = source if source else ""
                cells[1].text = target if target else ""
                cells[2].text = final_url
                cells[3].text = status

                # Color coding based on status (only works with some Word readers)
                run = cells[3].paragraphs[0].add_run(f" ({status})")
                if "Success" in status:
                    run.font.color.rgb = RGBColor(0, 128, 0)  # Green
                elif "Mismatch" in status:
                    run.font.color.rgb = RGBColor(255, 0, 0)  # Red
                else:
                    run.font.color.rgb = RGBColor(255, 165, 0)  # Orange

                self.progress["value"] = i - 1
                self.status_label.config(text=f"Checking row {i}/{total}")
                time.sleep(0.5)

            output_file = self.filepath.replace(".xlsx", "_report.docx")
            document.save(output_file)
            self.status_label.config(text="Completed! Saved as: " + output_file)
            messagebox.showinfo("Done", f"Report saved to:\n{output_file}")

        except Exception as e:
            messagebox.showerror("Error", str(e))

if __name__ == "__main__":
    root = tk.Tk()
    app = RedirectCheckerApp(root)
    root.mainloop()
